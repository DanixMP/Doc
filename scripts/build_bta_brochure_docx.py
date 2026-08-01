#!/usr/bin/env python3
"""Build two BTA brochure DOCX files from extracted PDF assets.

1) Exact replica — landscape A4 pages, each filled with the matching scanned page image.
2) Redesign — portrait RTL Persian catalog with structured text + related cropped images.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor, Twips
from PIL import Image, ImageDraw, ImageFont

ROOT = Path("/workspace")
PDF_EXT = ROOT / "extracted" / "pdf"
HIRES = PDF_EXT / "pages_hires"
RENDERS = PDF_EXT / "page_renders"
ASSETS = ROOT / "docs" / "bta-assets"
OUT_EXACT = ROOT / "docs" / "کاتالوگ-BTA-نسخه-عینی.docx"
OUT_REDESIGN = ROOT / "docs" / "کاتالوگ-BTA-نسخه-بازطراحی.docx"
FONT_DIR = ROOT / "docs" / "fonts"

FONT_REG = FONT_DIR / "Vazirmatn-Regular.ttf"
FONT_BOLD = FONT_DIR / "Vazirmatn-Bold.ttf"
FONT_MED = FONT_DIR / "Vazirmatn-Medium.ttf"
FONT_NAME = "Vazirmatn"

# BTA brand palette (black / gold / warm stone — matches brochure)
C_BLACK = "#0B0B0B"
C_INK = "#1A1A1A"
C_GOLD = "#C4A35A"
C_GOLD_DARK = "#9A7A3A"
C_STONE = "#F3EEE4"
C_MUTED = "#6B6560"
C_WHITE = "#FFFFFF"
C_LINE = "#E2D9C8"
C_PANEL = "#1F1C18"


def to_fa(s) -> str:
    return str(s).translate(str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹"))


def hex_rgb(h: str):
    h = h.lstrip("#")
    return tuple(int(h[i : i + 2], 16) for i in (0, 2, 4))


def load_font(path: Path, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(path), size=size)


def draw_rtl(draw, xy, text, font, fill, anchor="rt"):
    draw.text(xy, text, font=font, fill=fill, anchor=anchor, direction="rtl", language="fa")


def ensure_cs_rtl_on_run(run, size_pt: float, bold: bool = False, color: str | None = None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), FONT_NAME)

    for tag in ("w:rtl", "w:cs", "w:bCs", "w:sz", "w:szCs"):
        for old in rPr.findall(qn(tag)):
            rPr.remove(old)

    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)
    rPr.append(OxmlElement("w:cs"))
    if bold:
        rPr.append(OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(szCs)


def set_paragraph_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:bidi")):
        pPr.remove(old)
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def add_p(doc, text, size=11, bold=False, color=None, after=8, before=0,
          align=WD_ALIGN_PARAGRAPH.RIGHT):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, align=align)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    ensure_cs_rtl_on_run(run, size, bold=bold, color=color or C_INK)
    return p


def add_h(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    colors = {1: C_BLACK, 2: C_GOLD_DARK, 3: C_INK}
    p = add_p(doc, text, size=sizes[level], bold=True, color=colors[level], before=16, after=6)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "14")
        bottom.set(qn("w:space"), "5")
        bottom.set(qn("w:color"), C_GOLD.replace("#", ""))
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color.replace("#", ""))
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell(cell, text, bold=False, size=9.5, color=C_INK, fill=None, center=True):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    ensure_cs_rtl_on_run(run, size, bold=bold, color=color)


def set_table_rtl(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:bidiVisual")):
        tblPr.remove(old)
    bidi = OxmlElement("w:bidiVisual")
    bidi.set(qn("w:val"), "1")
    tblPr.append(bidi)


def add_table(doc, title, headers, rows, col_widths=None):
    add_p(doc, title, size=11, bold=True, color=C_BLACK, before=10, after=4)
    headers_full = ["ردیف"] + headers
    rows_full = [[to_fa(i + 1)] + list(row) for i, row in enumerate(rows)]
    table = doc.add_table(rows=1 + len(rows_full), cols=len(headers_full))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_rtl(table)
    for j, h in enumerate(headers_full):
        set_cell(table.rows[0].cells[j], h, bold=True, size=9.5, color=C_WHITE, fill=C_BLACK)
    for i, row in enumerate(rows_full):
        fill = C_STONE if i % 2 == 0 else C_WHITE
        for j, val in enumerate(row):
            set_cell(table.rows[i + 1].cells[j], str(val), size=9, fill=fill)
    if col_widths:
        widths = [1.2] + col_widths
        widths = (widths + [2.5] * len(headers_full))[: len(headers_full)]
        for row in table.rows:
            for idx, w in enumerate(widths):
                row.cells[idx].width = Cm(w)
    doc.add_paragraph()
    return table


def add_pic(doc, path, width_cm=16.5, caption: str | None = None):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    if caption:
        add_p(doc, caption, size=9, color=C_MUTED, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)


def set_doc_rtl_lang(doc: Document):
    settings = doc.settings.element
    for old in settings.findall(qn("w:themeFontLang")):
        settings.remove(old)
    lang = OxmlElement("w:themeFontLang")
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:eastAsia"), "en-US")
    lang.set(qn("w:bidi"), "fa-IR")
    settings.append(lang)


def set_section_bidi(section):
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:bidi")):
        sectPr.remove(old)
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    sectPr.append(bidi)


def add_page_number_footer(section, label="صفحه"):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(f"{label} ")
    ensure_cs_rtl_on_run(run, 9, color=C_MUTED)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run2 = p.add_run()
    run2._r.append(fld1)
    run2._r.append(instr)
    run2._r.append(fld2)
    ensure_cs_rtl_on_run(run2, 9, color=C_MUTED)


def disable_spacing(paragraph):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


# ---------------------------------------------------------------------------
# Asset preparation for redesign
# ---------------------------------------------------------------------------

def crop_save(src: Path, box, dest: Path, max_w=1400):
    """box = (left, top, right, bottom) in pixels of src."""
    im = Image.open(src).convert("RGB")
    crop = im.crop(box)
    if crop.width > max_w:
        ratio = max_w / crop.width
        crop = crop.resize((max_w, int(crop.height * ratio)), Image.Resampling.LANCZOS)
    dest.parent.mkdir(parents=True, exist_ok=True)
    crop.save(dest, quality=90, optimize=True)
    return dest


def prepare_redesign_assets():
    ASSETS.mkdir(parents=True, exist_ok=True)

    # Cover banner — black/gold brand composition using cover photo crop
    cover_src = HIRES / "page_01.jpg"
    im = Image.open(cover_src).convert("RGB")
    w, h = im.size
    # left building photo region
    photo = im.crop((0, 0, int(w * 0.55), int(h * 0.72))).resize((1100, 720), Image.Resampling.LANCZOS)

    banner = Image.new("RGB", (1800, 980), hex_rgb(C_BLACK))
    banner.paste(photo, (0, 0))
    # gold diagonal wedge
    overlay = Image.new("RGBA", banner.size, (0, 0, 0, 0))
    od = ImageDraw.Draw(overlay)
    od.polygon([(980, 0), (1800, 0), (1800, 980), (720, 980)], fill=(*hex_rgb(C_BLACK), 235))
    od.polygon([(980, 0), (1040, 0), (780, 980), (720, 980)], fill=(*hex_rgb(C_GOLD), 255))
    banner = Image.alpha_composite(banner.convert("RGBA"), overlay).convert("RGB")
    draw = ImageDraw.Draw(banner)
    draw_rtl(draw, (1680, 220), "بهین تجربه آرتیمان", load_font(FONT_BOLD, 64), hex_rgb(C_GOLD), "rt")
    draw_rtl(draw, (1680, 300), "شرکت هلدینگ بتا", load_font(FONT_MED, 34), (230, 220, 200), "rt")
    draw.rectangle([1380, 340, 1680, 346], fill=hex_rgb(C_GOLD))
    services = ["طراحی", "محاسبه", "تولید", "اجرا", "تعمیر و نگهداری"]
    y = 400
    for s in services:
        draw_rtl(draw, (1680, y), f"■  {s}", load_font(FONT_REG, 30), (245, 240, 230), "rt")
        y += 55
    draw_rtl(
        draw, (1680, 760),
        "انواع پروژه‌های ساختمانی، سوله و تهویه مطبوع خانگی و صنعتی",
        load_font(FONT_MED, 24), hex_rgb(C_GOLD), "rt",
    )
    draw_rtl(draw, (1680, 880), "BTA  ·  کاتالوگ محصولات و خدمات", load_font(FONT_REG, 26), (200, 190, 175), "rt")
    banner.save(ASSETS / "cover_banner.jpg", quality=92)

    # Related photo crops (approximate regions from hi-res landscape pages)
    crops = {
        "intro_equipment.jpg": (HIRES / "page_02.jpg", (40, 120, 1450, 1750)),
        "products_overview.jpg": (HIRES / "page_05.jpg", (80, 80, 1600, 1750)),
        "cooling_tower.jpg": (HIRES / "page_06.jpg", (80, 900, 1300, 1750)),
        "ahu_truck.jpg": (HIRES / "page_10.jpg", (40, 80, 1750, 1750)),
        "hygienic_ahu.jpg": (HIRES / "page_12.jpg", (40, 80, 1400, 1100)),
        "filters.jpg": (HIRES / "page_13.jpg", (40, 200, 1200, 1700)),
        "laminar_or.jpg": (HIRES / "page_14.jpg", (80, 200, 1300, 1600)),
        "unit_heater.jpg": (HIRES / "page_19.jpg", (40, 80, 900, 900)),
        "projects_towers.jpg": (HIRES / "page_22.jpg", (500, 80, 1700, 1750)),
        "warehouse.jpg": (HIRES / "page_26.jpg", (40, 80, 1000, 1750)),
        "diadaru_gallery.jpg": (HIRES / "page_27.jpg", (200, 80, 2500, 1750)),
        "cleanroom.jpg": (HIRES / "page_28.jpg", (900, 900, 2500, 1700)),
        "contact_skyline.jpg": (HIRES / "page_32.jpg", (40, 1100, 900, 1750)),
    }
    for name, (src, box) in crops.items():
        if src.exists():
            # clamp box to image size
            im = Image.open(src)
            l, t, r, b = box
            r = min(r, im.width)
            b = min(b, im.height)
            l = max(0, min(l, r - 10))
            t = max(0, min(t, b - 10))
            crop_save(src, (l, t, r, b), ASSETS / name)

    # Section divider strip
    strip = Image.new("RGB", (1600, 120), hex_rgb(C_BLACK))
    sd = ImageDraw.Draw(strip)
    sd.rectangle([0, 100, 1600, 120], fill=hex_rgb(C_GOLD))
    draw_rtl(sd, (1520, 48), "بهین تجربه آرتیمان  ·  BTA", load_font(FONT_MED, 36), hex_rgb(C_GOLD), "rt")
    strip.save(ASSETS / "section_strip.jpg", quality=90)

    # Contact footer graphic
    contact = Image.new("RGB", (1600, 520), hex_rgb(C_BLACK))
    cd = ImageDraw.Draw(contact)
    cd.polygon([(0, 0), (420, 0), (280, 520), (0, 520)], fill=hex_rgb(C_GOLD))
    sky = ASSETS / "contact_skyline.jpg"
    if sky.exists():
        sk = Image.open(sky).convert("RGB").resize((380, 480), Image.Resampling.LANCZOS)
        contact.paste(sk, (20, 20))
    draw_rtl(cd, (1520, 80), "ارتباط با ما", load_font(FONT_BOLD, 48), hex_rgb(C_GOLD), "rt")
    draw_rtl(cd, (1520, 160), "دفتر مرکزی: تهران، پاسداران، میدان بنی‌هاشم", load_font(FONT_REG, 26), (235, 230, 220), "rt")
    draw_rtl(cd, (1520, 210), "مجتمع لاله‌سنتر، طبقه سوم، واحد ۳۰۱", load_font(FONT_REG, 26), (235, 230, 220), "rt")
    draw_rtl(cd, (1520, 280), "تلفن: ۲۶۱۴۷۵۲۹ - ۲۶۱۴۷۳۲۳ - ۲۶۱۴۷۳۲۱ - ۲۶۱۴۷۳۱۱ (۰۲۱)", load_font(FONT_MED, 24), hex_rgb(C_GOLD), "rt")
    draw_rtl(cd, (1520, 330), "همراه: ۰۹۱۹۹۵۲۲۲۷۱", load_font(FONT_MED, 24), hex_rgb(C_GOLD), "rt")
    draw_rtl(cd, (1520, 400), "www.btamep.com   ·   btamep@gmail.com", load_font(FONT_REG, 24), (210, 200, 185), "rt")
    contact.save(ASSETS / "contact_panel.jpg", quality=92)

    print(f"Prepared assets in {ASSETS} ({len(list(ASSETS.glob('*')))} files)")


# ---------------------------------------------------------------------------
# 1) Exact replica DOCX
# ---------------------------------------------------------------------------

def _configure_landscape_zero_margin(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)


def _put_fullpage_image(paragraph, img_path: Path, width_cm=29.7):
    """Clear paragraph and insert a single full-width page image."""
    # Remove any existing runs
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    disable_spacing(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))


def build_exact_replica():
    """One landscape Word page per PDF page, image edge-to-edge."""
    doc = Document()
    set_doc_rtl_lang(doc)

    _configure_landscape_zero_margin(doc.sections[0])
    usable_w = 29.7  # A4 landscape width in cm

    for i in range(1, 33):
        img_path = HIRES / f"page_{i:02d}.jpg"
        if not img_path.exists():
            img_path = RENDERS / f"page_{i:02d}.png"

        if i == 1:
            # python-docx may start with zero or one empty paragraph
            p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        else:
            # add_section creates a new empty paragraph — reuse it (no extra blank para)
            new_sec = doc.add_section()
            _configure_landscape_zero_margin(new_sec)
            p = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()

        _put_fullpage_image(p, img_path, width_cm=usable_w)
        print(f"exact: page {i:02d} <- {img_path.name}")

    OUT_EXACT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_EXACT)
    print(f"Wrote {OUT_EXACT} ({OUT_EXACT.stat().st_size // 1024} KB)")
    return OUT_EXACT


# ---------------------------------------------------------------------------
# 2) Redesigned DOCX
# ---------------------------------------------------------------------------

def build_redesign():
    prepare_redesign_assets()
    doc = Document()
    set_doc_rtl_lang(doc)

    section = doc.sections[0]
    set_section_bidi(section)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.right_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.6)
    add_page_number_footer(section)

    # ---- Cover ----
    add_pic(doc, ASSETS / "cover_banner.jpg", width_cm=17.4)
    add_p(
        doc,
        "کاتالوگ محصولات، خدمات و پروژه‌های اجراشده",
        size=13, bold=True, color=C_GOLD_DARK,
        align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=4,
    )
    add_p(
        doc,
        "نسخه بازطراحی‌شده بر اساس کاتالوگ چاپی شرکت",
        size=10, color=C_MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER, after=18,
    )

    # ---- Intro ----
    add_h(doc, "مقدمه", 1)
    add_pic(doc, ASSETS / "intro_equipment.jpg", 16.5, "تجهیزات سرمایش و تهویه مطبوع صنعتی")
    add_p(
        doc,
        "شرکت ساختمانی و تأمین تجهیزات بهین تجربه آرتیمان (BTA) در سال ۱۳۹۹ با هدف طراحی و اجرای "
        "انواع پروژه‌های ساختمانی، سوله و سیستم‌های تهویه مطبوع از صفر تا صد تأسیس شده است.",
        size=11, after=8,
    )
    add_p(
        doc,
        "حوزه فعالیت شامل پروژه‌های اداری، تجاری، آموزشی، فرهنگی و نظامی است؛ از جمله انبارهای دارویی "
        "و نگهداری مواد غذایی. شرکت بر به‌کارگیری تجهیزات دارای گواهی‌های معتبر داخلی و بین‌المللی "
        "و آمادگی برای پروژه‌های بین‌المللی تأکید دارد.",
        size=11, after=10,
    )

    # ---- Capabilities ----
    add_h(doc, "توانمندی‌ها: تولیدات و اجرایی‌ات", 1)
    add_h(doc, "تولیدات", 2)
    for item in [
        "طراحی و تولید انواع چیلرهای تراکمی آب‌خنک و هواخنک",
        "طراحی و تولید انواع برج‌های خنک‌کننده",
        "طراحی و تولید انواع هواساز، هواسازهای هایژنیک و ایرواشر",
        "طراحی و تولید انواع روفتاپ پکیج (پکیج یونیت)",
        "طراحی و تولید انواع لامینار باکس",
        "طراحی و تولید انواع فیلتر هپا",
        "طراحی و تولید انواع بویلر صنعتی، کویل و فن‌کویل",
        "طراحی، تولید، اجرا، تعمیر و نگهداری انواع سیستم‌های تهویه مطبوع، داکت اسپلیت و کولر",
    ]:
        add_p(doc, f"• {item}", size=10.5, after=3)

    add_h(doc, "اجرایی‌ات", 2)
    for item in [
        "طراحی و اجرای صفر تا صد انواع پروژه‌های ساختمانی",
        "طراحی، محاسبه، اجرا و تجهیز انواع سوله‌های صنعتی",
        "طراحی و اجرای انواع کلین‌روم",
        "طراحی، اجرا، تأمین، تجهیز و نگهداری سیستم‌های موتورخانه مکانیکال",
        "طراحی و اجرای سیستم‌های برق قدرت، شبکه روشنایی و دوربین مداربسته",
    ]:
        add_p(doc, f"• {item}", size=10.5, after=3)

    add_pic(doc, ASSETS / "products_overview.jpg", 16.5, "نمونه‌ای از محصولات تهویه و سرمایش")

    # ---- Chillers ----
    add_h(doc, "چیلرهای تراکمی", 1)
    add_p(
        doc,
        "طراحی سفارشی بر اساس نیاز مشتری و شرایط اقلیمی، با کندانسور آب‌خنک یا هواخنک؛ "
        "کنترل هوشمند مبتنی بر پی‌ال‌سی و قابلیت اتصال به سیستم مدیریت ساختمان؛ "
        "اواپراتور پوسته‌ولوله با لوله‌های مسی؛ بدنه گالوانیزه با پوشش پودری الکترواستاتیک؛ "
        "درجه حفاظت آی‌پی ۵۴.",
        size=11, after=8,
    )
    add_table(
        doc,
        "جدول ۱. خطوط محصول چیلر تراکمی",
        ["نوع", "کمپرسور", "حدود ظرفیت (تن)"],
        [
            ["اسکرول", "هرمتیک", "۳٫۵ تا ۱۲۰"],
            ["پیستونی", "نیمه‌هرمتیک", "۳٫۵ تا ۱۲۰"],
            ["اسکرو", "نیمه‌هرمتیک", "۵۰ تا ۱۲۸۰"],
        ],
        col_widths=[3.5, 3.5, 4],
    )
    add_p(doc, "برندهای کمپرسور: کوپلند، بیتزر، دانفوس، هانبل", size=10, color=C_MUTED, after=10)

    # ---- Cooling towers ----
    add_h(doc, "برج‌های خنک‌کننده", 1)
    add_pic(doc, ASSETS / "cooling_tower.jpg", 14, "برج خنک‌کننده فایبرگلاس / مکعبی")
    add_p(
        doc,
        "برج خنک‌کننده برای خنک‌کاری آب فرآیندی در مجتمع‌های مسکونی، تجاری و صنعتی "
        "(پالایشگاه، نیروگاه، کارخانه) از طریق تبادل حرارت آب و هوا به‌کار می‌رود. "
        "ظرفیت‌های متداول از حدود ۵ تا بیش از ۱۲۵۰ تن و در مدل‌های جریان متقاطع و جریان مخالف ارائه می‌شود.",
        size=11, after=6,
    )
    add_h(doc, "انواع برج", 2)
    for item in [
        "برج گرد یا بطری‌شکل جریان مخالف — ظرفیت حدود ۸ تا ۱۴۰۰ تن",
        "برج فایبرگلاس جریان متقاطع — ظرفیت حدود ۵۰ تا ۱۰۰۰ تن؛ صدای کم و فضای نصب محدود",
        "برج مکعب جریان مخالف — مدار باز و مدار بسته؛ پکینگ یو‌پی‌وی‌سی",
        "برج مدار بسته — جلوگیری از رسوب و خوردگی با جدا نگه داشتن آب فرآیندی در حلقه بسته",
    ]:
        add_p(doc, f"• {item}", size=10.5, after=3)

    # ---- AHU ----
    add_h(doc, "هواساز و تجهیزات مرتبط", 1)
    add_pic(doc, ASSETS / "ahu_truck.jpg", 16.5, "هواساز صنعتی مدولار")
    add_p(
        doc,
        "هواساز دستگاه تأمین هوای مطبوع برای کنترل دما و رطوبت است و از کویل‌های گرمایش و سرمایش "
        "به‌همراه فیلتراسیون ذرات معلق (آلومینیومی، کیسه‌ای، کربن اکتیو، هپا و اولپا) بهره می‌برد. "
        "ظرفیت متداول حدود ۱۵۰۰ تا ۶۵۰۰۰ سی‌اف‌ام؛ طراحی تک‌زون یا چندزون؛ مدل افقی یا عمودی.",
        size=11, after=8,
    )
    add_pic(doc, ASSETS / "hygienic_ahu.jpg", 15, "هواساز هایژنیک")
    add_h(doc, "هواساز هایژنیک", 2)
    add_p(
        doc,
        "برای محیط‌های حساس مانند بیمارستان، تولید دارو و واکسن، صنایع غذایی و اتاق تمیز. "
        "ورق استنلس استیل، گوشه‌های گرد، فن کوپل مستقیم، و مستندات اعتبارسنجی "
        "آی‌کیو، دی‌کیو، او‌کیو و پی‌کیو. استانداردهای مرجع: ایزو ۱۴۶۴۴ و وی‌دی‌آی ۶۰۲۲.",
        size=11, after=6,
    )
    add_h(doc, "فیلترها", 2)
    add_pic(doc, ASSETS / "filters.jpg", 12, "انواع فیلتر فلت، کیسه‌ای، هپا و پنلی")
    for item in [
        "فیلتر فلت پلی‌استر کلاس جی۳ (و آلومینیومی قابل شستشو کلاس جی۲ بنا به درخواست)",
        "فیلتر کیسه‌ای کلاس ام۵ تا اف۹ — طول استاندارد ۳۸۰ میلی‌متر",
        "فیلتر هپا کلاس اچ۱۳ و اچ۱۴ مطابق ای‌ان ۱۸۲۲",
        "فیلتر پنلی چین‌دار به‌عنوان پیش‌فیلتر",
    ]:
        add_p(doc, f"• {item}", size=10.5, after=3)

    # ---- Laminar ----
    add_h(doc, "لامینار باکس و جریان آرام", 1)
    add_pic(doc, ASSETS / "laminar_or.jpg", 15, "محیط استریل / اتاق عمل")
    add_p(
        doc,
        "سیستم جریان هوای آرام برای جلوگیری از آلودگی تجهیزات و محیط با ایجاد جریان یکنواخت "
        "هوای فیلترشده. انواع سقفی، عمودی، افقی و کابینتی. لامینار سقفی اتاق عمل با فیلتر هپا "
        "اچ۱۳/اچ۱۴ و سرعت جریان حدود ۰٫۲ تا ۰٫۳ متر بر ثانیه.",
        size=11, after=6,
    )
    add_p(doc, "کاربردها: اتاق عمل، تولید دارو و واکسن، لیزر، اپتیک، میکروالکترونیک، بسته‌بندی غذایی و دارویی، مونتاژ.", size=10.5, after=10)

    # ---- Fan coil / heaters ----
    add_h(doc, "فن‌کوئل، یونیت هیتر و سایر محصولات", 1)
    add_h(doc, "فن‌کوئل", 2)
    add_p(
        doc,
        "کویل سه‌ردیفه با هدر ثابت؛ فن ای‌بی‌اس خودبالانس؛ گواهی سطح صدا از آزمایشگاه‌های معتبر. "
        "داکت فن‌کوئل با کابینت گالوانیزه ۱ میلی‌متری، کویل ۴ ردیفه مسی، سینی درین اپوکسی، "
        "فیلتر آلومینیومی قابل شستشو و امکان کویل برقی.",
        size=11, after=6,
    )
    add_h(doc, "یونیت هیتر", 2)
    add_pic(doc, ASSETS / "unit_heater.jpg", 10, "یونیت هیتر صنعتی")
    add_p(
        doc,
        "گرمایش فضاهای بزرگ با سقف بلند (سالن ورزشی، سوله، استخر). مدل‌های آبگرم، بخار و روغن داغ. "
        "ظرفیت تقریبی از ۴۲٬۰۰۰ تا ۴۵۰٬۰۰۰ بی‌تی‌یو بر ساعت.",
        size=11, after=6,
    )
    add_table(
        doc,
        "جدول ۲. مشخصات تقریبی یونیت هیتر آبگرم",
        ["مدل", "ظرفیت (بی‌تی‌یو)", "هوادهی (مترمکعب/ساعت)", "توان فن (وات)"],
        [
            ["اچ‌دبلیو-۵۰", to_fa(50000), to_fa(2000), to_fa(110)],
            ["اچ‌دبلیو-۸۰", to_fa(80000), to_fa(2700), to_fa(155)],
            ["اچ‌دبلیو-۱۲۰", to_fa(120000), to_fa(3300), to_fa(240)],
            ["اچ‌دبلیو-۱۸۰", to_fa(180000), to_fa(5600), to_fa(300)],
            ["اچ‌دبلیو-۲۰۰", to_fa(200000), to_fa(5600), to_fa(300)],
            ["اچ‌دبلیو-۲۵۰", to_fa(250000), to_fa(5600), to_fa(300)],
        ],
        col_widths=[2.8, 3.2, 4, 2.8],
    )
    add_h(doc, "دیگر محصولات", 2)
    for item in [
        "زنت دوفصلی — سرمایش تبخیری و گرمایش با کویل آب گرم؛ حدود ۱۰۰۰ تا ۲۶۰۰۰ سی‌اف‌ام",
        "فن سانتریفیوژ برای هواساز، اگزاست پارکینگ و محیط صنعتی",
        "پرده هوا — طول حدود ۹۰ سانتی‌متر تا ۲ متر",
        "بویلر فولادی — حدود ۱۰۰٬۰۰۰ تا ۱٬۰۰۰٬۰۰۰ کیلوکالری بر ساعت (برخی مدل‌ها بالاتر)",
        "هیتر تابشی سقفی و هیتر برقی فن‌دار صنعتی",
        "ایرواشر و کویل‌های صنعتی با پوشش گلد فین / بلو فین",
    ]:
        add_p(doc, f"• {item}", size=10.5, after=3)

    # ---- Cleanroom ----
    add_h(doc, "اتاق تمیز", 1)
    add_pic(doc, ASSETS / "cleanroom.jpg", 15, "نمونه فضای اتاق تمیز")
    add_p(
        doc,
        "اتاق تمیز محیطی کنترل‌شده برای تولید یا تحقیقات علمی و صنعتی است که مقدار آلاینده‌های "
        "زیست‌محیطی (گردوغبار، میکروب معلق، بخار) در آن بسیار پایین‌تر از فضای بسته معمول نگه داشته می‌شود. "
        "طبقه‌بندی بر اساس استانداردهای ایزو ۱۴۶۴۴-۱، فدرال ۲۰۹ئی و بی‌اس ۵۲۹۵ انجام می‌شود.",
        size=11, after=6,
    )
    add_table(
        doc,
        "جدول ۳. معادل‌های متداول ایزو ۱۴۶۴۴-۱",
        ["کلاس ایزو", "معادل تقریبی فدرال ۲۰۹ئی", "ذرات (≥۰٫۵ میکرون) در مترمکعب"],
        [
            ["ایزو ۵", "کلاس ۱۰۰", to_fa(3520)],
            ["ایزو ۷", "کلاس ۱۰٬۰۰۰", to_fa(352000)],
            ["ایزو ۹", "هوای اتاق", to_fa(35200000)],
        ],
        col_widths=[3, 4, 5],
    )
    add_p(
        doc,
        "کاربردها: داروسازی، تجهیزات پزشکی، صنایع غذایی، هوافضا، سامانه‌های دفاعی، میکروالکترونیک، هسته‌ای و کشاورزی.",
        size=10.5, after=10,
    )

    # ---- Projects ----
    add_h(doc, "نمونه‌ای از پروژه‌های اجراشده", 1)
    add_pic(doc, ASSETS / "projects_towers.jpg", 15, "نصب برج‌های خنک‌کننده در پروژه‌های مختلف")
    add_h(doc, "برج‌های خنک‌کننده", 2)
    for item in [
        "آتی‌ساز — برج ۸۰۰ تن (دو دستگاه)",
        "پرتوآبی شهرک صنعتی بهارستان — برج ۴۰۰ تن",
        "چاپخانه ارتش میدان سبلان — برج ۲۰۰ تن",
        "پارکینگ طبقاتی چیتگر",
        "شرکت نماد جاده مخصوص — برج ۲۵۰ تن",
        "فروشگاه شهروند؛ نانوی ری‌دی عباس‌آباد؛ مهندسی گاز ۱۰۰ تن؛ بیمارستان آپادانا ۸۰ تن",
    ]:
        add_p(doc, f"• {item}", size=10.5, after=3)

    add_h(doc, "داروسازی و بیمارستان", 2)
    add_p(doc, "رحمان دارو، کیمیدارو، داروسازی سها، دماوند دارو، دایا دارو", size=10.5, after=3)
    add_p(
        doc,
        "بیمارستان‌ها: کودکان تبریز، مریم البرز، امام خمینی استهبان، مدنی تبریز، حکیم فارابی، شهریار تبریز، شریعتی تهران",
        size=10.5, after=6,
    )

    add_h(doc, "سازمانی، بانکی و داکت اسپلیت", 2)
    for item in [
        "صنایع دفاعی وزارت دفاع — اجرای صفر تا صد داکت اسپلیت",
        "بانک تجارت (مطهری، ابن‌سینا، اداره مرکزی) — تجهیز کولرگازی",
        "اداره مرکزی بانک ملی شعبه غرب تهران — تجهیز، تعمیر و نگهداری تهویه",
        "دانشگاه امام صادق (ع) و پروژه‌های مسکونی در فشم، سوهانک، پاسداران، هروی، نواب و فتح",
    ]:
        add_p(doc, f"• {item}", size=10.5, after=3)

    add_h(doc, "ساخت و تجهیز سوله", 2)
    add_pic(doc, ASSETS / "warehouse.jpg", 12, "نمونه فضای انبار تجهیزشده")
    add_p(
        doc,
        "بیش از یک دهه تجربه در طراحی، محاسبه و تجهیز سوله صنعتی؛ بیش از ۵۰۰ پروژه موفق؛ "
        "فضاهای بزرگ بیش از ۶۰۰۰ مترمربع؛ تمرکز بر انبارهای دارویی، غذایی و پوشاک.",
        size=11, after=6,
    )
    add_pic(doc, ASSETS / "diadaru_gallery.jpg", 16.5, "نمونه تجهیز سوله‌های دیادارو")
    add_p(
        doc,
        "همکاران نمونه: دایادارو / دانیادارو، بانک تجارت، بانک صادرات، بانک ایران‌زمین، بانک ملت، "
        "سازمان تأمین اجتماعی، سازمان صداوسیما.",
        size=10.5, after=10,
    )

    # ---- Maintenance ----
    add_h(doc, "تعمیر و نگهداری موتورخانه", 1)
    add_p(
        doc,
        "سابقه حدود ۱۰ ساله در تجهیز، تعمیر، مراقبت و نگهداری سیستم‌های موتورخانه. "
        "از جمله: مرکز ملی مخابرات (سرمایش دیتاسنتر)، برج ارک، برج‌های مهستان بلوک ای۲/ای۳، "
        "پژوهشگاه علوم انسانی (چند ساختمان)، ساختمان رویال، برج چهارباغ و ساختمان آرش.",
        size=11, after=12,
    )

    # ---- Contact ----
    add_h(doc, "ارتباط با ما", 1)
    add_pic(doc, ASSETS / "contact_panel.jpg", 17.4)
    add_p(
        doc,
        "Unit 301, Floor 3, Laleh Center Complex, Bani Hashem Ave, Pasdaran, Tehran",
        size=9.5, color=C_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=4,
    )
    add_p(
        doc,
        "www.btamep.com  ·  btamep@gmail.com",
        size=10, bold=True, color=C_GOLD_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=8,
    )

    OUT_REDESIGN.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_REDESIGN)
    print(f"Wrote {OUT_REDESIGN} ({OUT_REDESIGN.stat().st_size // 1024} KB)")
    return OUT_REDESIGN


def main():
    print("=== Exact replica ===")
    build_exact_replica()
    print("=== Redesign ===")
    build_redesign()
    print("Done.")


if __name__ == "__main__":
    main()
