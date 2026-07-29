#!/usr/bin/env python3
"""Generate a beautiful RTL Persian proposal DOCX with Vazirmatn font."""

from __future__ import annotations

import io
import shutil
import zipfile
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from matplotlib import font_manager
from matplotlib.patches import FancyBboxPatch
from PIL import Image, ImageDraw, ImageFont, ImageFilter

ROOT = Path("/workspace")
FONT_DIR = ROOT / "fonts" / "vazirmatn" / "fonts" / "ttf"
ASSETS = ROOT / "docs" / "assets"
DOC_FONTS = ROOT / "docs" / "fonts"
OUT_DOCX = ROOT / "docs" / "پیشنهادیه-سامانه-هوشمند-پاپ-اسمیر.docx"

FONT_REG = FONT_DIR / "Vazirmatn-Regular.ttf"
FONT_BOLD = FONT_DIR / "Vazirmatn-Bold.ttf"
FONT_MED = FONT_DIR / "Vazirmatn-Medium.ttf"
FONT_LIGHT = FONT_DIR / "Vazirmatn-Light.ttf"
FONT_NAME = "Vazirmatn"

C_TEAL = "#0F6B6B"
C_TEAL_DARK = "#0A4F4F"
C_TEAL_LIGHT = "#E6F3F3"
C_CORAL = "#C45C4A"
C_GOLD = "#C4A35A"
C_INK = "#1C2B2B"
C_MUTED = "#5A6B6B"
C_BG = "#F7FAFA"
C_WHITE = "#FFFFFF"
C_LINE = "#D5E3E3"


def to_persian_digits(s) -> str:
    return str(s).translate(str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹"))


def load_font(path: Path, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(path), size=size)


def draw_rtl(draw: ImageDraw.ImageDraw, xy, text: str, font, fill, anchor="rt"):
    """Draw properly shaped Persian text (requires libraqm)."""
    draw.text(xy, text, font=font, fill=fill, anchor=anchor, direction="rtl", language="fa")


def hex_to_rgb(h: str):
    h = h.lstrip("#")
    return tuple(int(h[i:i + 2], 16) for i in (0, 2, 4))


# ---------------------------------------------------------------------------
# Cover + charts (Pillow for correct Persian shaping)
# ---------------------------------------------------------------------------

def make_cover_banner(path: Path):
    W, H = 1800, 980
    img = Image.new("RGB", (W, H), hex_to_rgb(C_TEAL_DARK))
    draw = ImageDraw.Draw(img)

    # vertical soft gradient
    base = hex_to_rgb(C_TEAL_DARK)
    top = (18, 95, 95)
    for y in range(H):
        t = y / H
        col = tuple(int(base[i] * (1 - t) + top[i] * t) for i in range(3))
        draw.line([(0, y), (W, y)], fill=col)

    overlay = Image.new("RGBA", (W, H), (0, 0, 0, 0))
    od = ImageDraw.Draw(overlay)
    od.ellipse([-180, -120, 520, 520], fill=(255, 255, 255, 30))
    od.ellipse([1250, 420, 1900, 1100], fill=(196, 163, 90, 55))
    od.ellipse([980, -160, 1500, 320], fill=(255, 255, 255, 22))
    od.rounded_rectangle([90, 120, 1710, 860], radius=36, outline=(255, 255, 255, 55), width=2)
    img = Image.alpha_composite(img.convert("RGBA"), overlay).convert("RGB")
    draw = ImageDraw.Draw(img)

    f_sm = load_font(FONT_MED, 36)
    f_md = load_font(FONT_BOLD, 44)
    f_lg = load_font(FONT_BOLD, 52)
    f_xl = load_font(FONT_BOLD, 40)

    draw_rtl(draw, (W // 2, 175), "بسمه تعالی", f_md, hex_to_rgb(C_GOLD), anchor="mm")
    draw_rtl(draw, (W // 2, 245), "فرم درخواست طرح توسعه فناوری", f_sm, (230, 242, 242), anchor="mm")

    # gold divider
    draw.rectangle([620, 280, 1180, 286], fill=hex_to_rgb(C_GOLD))

    draw_rtl(draw, (W // 2, 340), "پیشنهادیه جامع طرح", f_lg, (255, 255, 255), anchor="mm")

    title = (
        "توسعه سامانه هوشمند شناسایی و طبقه‌بندی ناهنجاری‌های\n"
        "سیتولوژیک در نمونه‌های پاپ اسمیر به‌منظور تشخیص\n"
        "سرطان دهانه رحم در بانوان با کمک هوش مصنوعی"
    )
    # multiline RTL
    y = 430
    for line in title.split("\n"):
        draw_rtl(draw, (W // 2, y), line, f_xl, (240, 250, 250), anchor="mm")
        y += 58

    draw.rectangle([0, H - 22, W, H], fill=hex_to_rgb(C_GOLD))
    draw_rtl(
        draw,
        (W // 2, H - 70),
        "آموزش محلی · محرمانگی داده · پشته فشرده · ۱۲ ماه",
        load_font(FONT_MED, 30),
        (220, 235, 235),
        anchor="mm",
    )
    img.save(path, quality=95)


def chart_timeline(path: Path):
    W, H = 1600, 720
    img = Image.new("RGB", (W, H), hex_to_rgb(C_BG))
    draw = ImageDraw.Draw(img)
    title_f = load_font(FONT_BOLD, 36)
    label_f = load_font(FONT_MED, 26)
    small_f = load_font(FONT_REG, 22)

    draw_rtl(draw, (W - 60, 45), "زمان‌بندی ۱۲ ماهه طرح", title_f, hex_to_rgb(C_TEAL_DARK), anchor="rt")

    phases = [
        ("راه‌اندازی و اخلاق", 0, 2, C_TEAL),
        ("جمع‌آوری و برچسب‌گذاری", 1.5, 5, C_TEAL_DARK),
        ("آموزش مدل‌ها", 3.5, 8, C_CORAL),
        ("نمونه CDS", 6.5, 10, C_GOLD),
        ("پایلوت بیمارستانی", 8.5, 11, "#3D8B8B"),
        ("مستندسازی و تجاری‌سازی", 10.5, 12, C_MUTED),
    ]
    left, right = 80, 1180
    top, row_h = 110, 80
    scale = (right - left) / 12

    # month grid
    for m in range(0, 13):
        x = left + m * scale
        draw.line([(x, top - 10), (x, top + len(phases) * row_h)], fill=hex_to_rgb(C_LINE), width=1)
        draw_rtl(draw, (x, top + len(phases) * row_h + 18), to_persian_digits(m), small_f,
                 hex_to_rgb(C_MUTED), anchor="mm")
    draw_rtl(draw, ((left + right) / 2, H - 40), "ماه", small_f, hex_to_rgb(C_MUTED), anchor="mm")

    for i, (name, start, end, color) in enumerate(phases):
        y0 = top + i * row_h + 18
        x0 = left + start * scale
        x1 = left + end * scale
        draw.rounded_rectangle([x0, y0, x1, y0 + 44], radius=10, fill=hex_to_rgb(color))
        draw_rtl(draw, (right + 30, y0 + 22), name, label_f, hex_to_rgb(C_INK), anchor="rt")

    img.save(path, quality=95)


def chart_cost_compare(path: Path):
    W, H = 1400, 780
    img = Image.new("RGB", (W, H), hex_to_rgb(C_BG))
    draw = ImageDraw.Draw(img)
    title_f = load_font(FONT_BOLD, 34)
    label_f = load_font(FONT_MED, 24)
    small_f = load_font(FONT_REG, 22)

    draw_rtl(draw, (W - 50, 40), "مقایسه هزینه تقریبی آموزش محلی و آنلاین", title_f,
             hex_to_rgb(C_TEAL_DARK), anchor="rt")

    items = [
        ("پشته محلی\n(پیشنهادی)", 4, C_TEAL),
        ("Roboflow Core\n+ مصرف اعتبار", 6.5, C_GOLD),
        ("Roboflow\nEnterprise", 18, C_CORAL),
    ]
    chart_top, chart_bottom = 120, 560
    chart_left, gap, bar_w = 180, 120, 180
    max_v = 22
    # y axis
    for v in [0, 5, 10, 15, 20]:
        y = chart_bottom - (v / max_v) * (chart_bottom - chart_top)
        draw.line([(140, y), (1200, y)], fill=hex_to_rgb(C_LINE), width=1)
        draw_rtl(draw, (125, y), to_persian_digits(v), small_f, hex_to_rgb(C_MUTED), anchor="rm")

    for i, (name, val, color) in enumerate(items):
        x0 = chart_left + i * (bar_w + gap)
        h = (val / max_v) * (chart_bottom - chart_top)
        y0 = chart_bottom - h
        draw.rounded_rectangle([x0, y0, x0 + bar_w, chart_bottom], radius=14, fill=hex_to_rgb(color))
        draw_rtl(draw, (x0 + bar_w / 2, y0 - 28), to_persian_digits(f"≈{val} هزار دلار"), label_f,
                 hex_to_rgb(C_INK), anchor="mm")
        lines = name.split("\n")
        yy = chart_bottom + 28
        for line in lines:
            # keep English bits LTR visually; Persian with rtl
            if any("\u0600" <= ch <= "\u06FF" for ch in line):
                draw_rtl(draw, (x0 + bar_w / 2, yy), line, label_f, hex_to_rgb(C_INK), anchor="mm")
            else:
                draw.text((x0 + bar_w / 2, yy), line, font=label_f, fill=hex_to_rgb(C_INK), anchor="mm")
            yy += 32

    draw_rtl(draw, (W / 2, H - 35), "هزینه تقریبی سال اول", small_f, hex_to_rgb(C_MUTED), anchor="mm")
    img.save(path, quality=95)


def chart_stack_flow(path: Path):
    W, H = 1700, 900
    img = Image.new("RGB", (W, H), hex_to_rgb(C_BG))
    draw = ImageDraw.Draw(img)
    title_f = load_font(FONT_BOLD, 34)
    box_f = load_font(FONT_BOLD, 22)
    sub_f = load_font(FONT_REG, 18)
    note_f = load_font(FONT_MED, 22)

    draw_rtl(draw, (W - 50, 40), "نمودار پشته فناوری فشرده (محلی و قابل نگهداری)", title_f,
             hex_to_rgb(C_TEAL_DARK), anchor="rt")

    def box(x, y, w, h, fill, title, subtitle):
        draw.rounded_rectangle([x, y, x + w, y + h], radius=18, fill=hex_to_rgb(fill))
        draw_rtl(draw, (x + w / 2, y + h / 2 - 14), title, box_f, (255, 255, 255), anchor="mm")
        if subtitle:
            # mixed: draw Persian subtitle RTL
            draw_rtl(draw, (x + w / 2, y + h / 2 + 18), subtitle, sub_f, (235, 245, 245), anchor="mm")

    def arrow(x1, y1, x2, y2):
        draw.line([(x1, y1), (x2, y2)], fill=hex_to_rgb(C_MUTED), width=3)
        # simple arrow head
        draw.polygon([(x2, y2), (x2 - 12, y2 - 7), (x2 - 12, y2 + 7)], fill=hex_to_rgb(C_MUTED))

    # Flow LTR for technical pipeline readability, Persian labels inside
    box(60, 340, 200, 120, C_TEAL_DARK, "اسکنر لام", "تصویر WSI")
    box(320, 340, 210, 120, C_TEAL, "OpenSlide", "خواندن و تایل")
    box(600, 180, 210, 120, "#3D8B8B", "QuPath", "برچسب‌گذاری")
    box(600, 500, 210, 120, C_MUTED, "NAS محلی", "ذخیره امن")
    box(880, 340, 230, 120, C_CORAL, "PyTorch", "Seg + Class + Fusion")
    box(1180, 180, 210, 120, C_GOLD, "MLflow", "شاخص و نسخه")
    box(1180, 500, 210, 120, "#2F6F8F", "FastAPI", "CDS بیمارستانی")
    box(1460, 340, 180, 120, C_TEAL_DARK, "پاتولوژیست", "تأیید نهایی")

    arrow(260, 400, 315, 400)
    # to QuPath
    draw.line([(530, 400), (560, 400), (560, 240), (595, 240)], fill=hex_to_rgb(C_MUTED), width=3)
    draw.polygon([(595, 240), (583, 233), (583, 247)], fill=hex_to_rgb(C_MUTED))
    # to NAS
    draw.line([(530, 400), (560, 400), (560, 560), (595, 560)], fill=hex_to_rgb(C_MUTED), width=3)
    draw.polygon([(595, 560), (583, 553), (583, 567)], fill=hex_to_rgb(C_MUTED))
    # QuPath/NAS to PyTorch
    draw.line([(810, 240), (850, 240), (850, 400), (875, 400)], fill=hex_to_rgb(C_MUTED), width=3)
    draw.line([(810, 560), (850, 560), (850, 400)], fill=hex_to_rgb(C_MUTED), width=3)
    draw.polygon([(875, 400), (863, 393), (863, 407)], fill=hex_to_rgb(C_MUTED))
    # PyTorch to MLflow/FastAPI
    draw.line([(1110, 400), (1140, 400), (1140, 240), (1175, 240)], fill=hex_to_rgb(C_MUTED), width=3)
    draw.polygon([(1175, 240), (1163, 233), (1163, 247)], fill=hex_to_rgb(C_MUTED))
    draw.line([(1110, 400), (1140, 400), (1140, 560), (1175, 560)], fill=hex_to_rgb(C_MUTED), width=3)
    draw.polygon([(1175, 560), (1163, 553), (1163, 567)], fill=hex_to_rgb(C_MUTED))
    # to pathologist
    draw.line([(1390, 240), (1425, 240), (1425, 400), (1455, 400)], fill=hex_to_rgb(C_MUTED), width=3)
    draw.line([(1390, 560), (1425, 560), (1425, 400)], fill=hex_to_rgb(C_MUTED), width=3)
    draw.polygon([(1455, 400), (1443, 393), (1443, 407)], fill=hex_to_rgb(C_MUTED))

    # note bar
    draw.rounded_rectangle([120, 780, 1580, 860], radius=16, fill=hex_to_rgb(C_TEAL_LIGHT),
                           outline=hex_to_rgb(C_LINE), width=2)
    draw_rtl(
        draw,
        (W / 2, 820),
        "همه مراحل داخل شبکه خصوصی بیمارستان/دانشگاه — بدون ارسال داده بیمار به ابر خارجی",
        note_f,
        hex_to_rgb(C_TEAL_DARK),
        anchor="mm",
    )
    img.save(path, quality=95)


def chart_privacy_scores(path: Path):
    W, H = 1500, 780
    img = Image.new("RGB", (W, H), hex_to_rgb(C_BG))
    draw = ImageDraw.Draw(img)
    title_f = load_font(FONT_BOLD, 34)
    label_f = load_font(FONT_MED, 24)
    small_f = load_font(FONT_REG, 20)

    draw_rtl(draw, (W - 50, 40), "مقایسه کیفی آموزش محلی و آنلاین", title_f,
             hex_to_rgb(C_TEAL_DARK), anchor="rt")

    cats = ["حریم خصوصی", "تناسب با WSI", "ادغام بالینی", "کنترل هزینه", "پذیرش بیمارستانی"]
    local = [5, 5, 5, 4, 5]
    online = [2, 2, 1, 2, 2]

    left, right = 420, 1280
    top, row_h = 120, 100
    for i, cat in enumerate(cats):
        y = top + i * row_h
        draw_rtl(draw, (left - 30, y + 35), cat, label_f, hex_to_rgb(C_INK), anchor="rt")
        # grid
        for s in range(1, 6):
            x = left + s / 5 * (right - left)
            draw.line([(x, y), (x, y + 70)], fill=hex_to_rgb(C_LINE), width=1)
        # bars
        lw = local[i] / 5 * (right - left)
        ow = online[i] / 5 * (right - left)
        draw.rounded_rectangle([left, y + 8, left + lw, y + 32], radius=8, fill=hex_to_rgb(C_TEAL))
        draw.rounded_rectangle([left, y + 40, left + ow, y + 64], radius=8, fill=hex_to_rgb(C_CORAL))

    # legend
    draw.rounded_rectangle([80, 680, 160, 710], radius=6, fill=hex_to_rgb(C_TEAL))
    draw_rtl(draw, (180, 695), "پشته محلی", small_f, hex_to_rgb(C_INK), anchor="rm")
    draw.rounded_rectangle([320, 680, 400, 710], radius=6, fill=hex_to_rgb(C_CORAL))
    draw_rtl(draw, (420, 695), "پلتفرم آنلاین", small_f, hex_to_rgb(C_INK), anchor="rm")

    for s in range(1, 6):
        x = left + s / 5 * (right - left)
        draw_rtl(draw, (x, top + len(cats) * row_h + 10), to_persian_digits(s), small_f,
                 hex_to_rgb(C_MUTED), anchor="mm")
    draw_rtl(draw, ((left + right) / 2, H - 35), "امتیاز تناسب (۱ تا ۵)", small_f,
             hex_to_rgb(C_MUTED), anchor="mm")
    img.save(path, quality=95)


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def set_run_font(run, size=11, bold=False, color=None, font_name=FONT_NAME):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size * 2)))
    # remove old szCs
    for old in rPr.findall(qn("w:szCs")):
        rPr.remove(old)
    rPr.append(szCs)
    # complex script bold
    if bold:
        bCs = OxmlElement("w:bCs")
        for old in rPr.findall(qn("w:bCs")):
            rPr.remove(old)
        rPr.append(bCs)


def set_paragraph_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")


def add_rtl_paragraph(doc, text, size=11, bold=False, color=None, space_after=8, space_before=0,
                      align=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=None):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, align=align)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading_rtl(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 12}
    colors = {1: C_TEAL_DARK, 2: C_TEAL, 3: C_TEAL}
    p = add_rtl_paragraph(doc, text, size=sizes.get(level, 12), bold=True,
                          color=colors.get(level, C_TEAL), space_before=14, space_after=6)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), C_GOLD.replace("#", ""))
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color.replace("#", ""))
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_rtl(cell, text, bold=False, size=10, color=C_INK, fill=None, center=False):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade_cell(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_r = list(reversed(headers))
    rows_r = [list(reversed(r)) for r in rows]
    for j, h in enumerate(headers_r):
        set_cell_rtl(table.rows[0].cells[j], h, bold=True, size=10, color=C_WHITE,
                     fill=C_TEAL_DARK, center=True)
    for i, row in enumerate(rows_r):
        fill = C_TEAL_LIGHT if i % 2 == 0 else C_WHITE
        for j, val in enumerate(row):
            set_cell_rtl(table.rows[i + 1].cells[j], str(val), size=9.5, fill=fill, center=True)
    if col_widths:
        widths_r = list(reversed(col_widths))
        for row in table.rows:
            for idx, w in enumerate(widths_r):
                row.cells[idx].width = Cm(w)
    doc.add_paragraph()
    return table


def set_section_rtl(section):
    sectPr = section._sectPr
    # remove existing bidi if any then add
    for old in sectPr.findall(qn("w:bidi")):
        sectPr.remove(old)
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    sectPr.append(bidi)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.right_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.8)


def add_page_number_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("صفحه ")
    set_run_font(run, size=9, color=C_MUTED)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run2 = p.add_run()
    run2._r.append(fldChar1)
    run2._r.append(instr)
    run2._r.append(fldChar2)
    set_run_font(run2, size=9, color=C_MUTED)


def add_picture_centered(doc, path, width_cm=16):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    p.paragraph_format.space_after = Pt(10)
    return p


def add_horizontal_line(doc, color=C_GOLD):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color.replace("#", ""))
    pBdr.append(bottom)
    pPr.append(pBdr)


def package_fonts():
    DOC_FONTS.mkdir(parents=True, exist_ok=True)
    for f in [FONT_REG, FONT_BOLD, FONT_MED, FONT_LIGHT]:
        shutil.copy2(f, DOC_FONTS / f.name)
    # also keep a copy inside docs for distribution
    readme = DOC_FONTS / "README.txt"
    readme.write_text(
        "Install these Vazirmatn TTF files on your system so Microsoft Word/LibreOffice "
        "renders the Persian proposal with the intended typeface.\n"
        "Font family name in the DOCX: Vazirmatn\n",
        encoding="utf-8",
    )


def build_cover(doc, banner_path: Path):
    add_picture_centered(doc, banner_path, width_cm=17)
    add_rtl_paragraph(doc, " ", size=6, space_after=2)

    meta = [
        ("حوزه فناوری", "تجهیزات پزشکی، دارو و سلامت + ICT"),
        ("مدت اجرا", "۱۲ ماه"),
        ("خروجی اصلی", "نمونه محصول + مجموعه‌داده بومی"),
        ("رویکرد فنی", "آموزش و استقرار کاملاً محلی (On-premise)"),
        ("پشته فشرده", "OpenSlide · QuPath · PyTorch · MLflow · FastAPI"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        set_cell_rtl(table.rows[i].cells[1], k, bold=True, size=10, color=C_WHITE, fill=C_TEAL, center=True)
        set_cell_rtl(table.rows[i].cells[0], v, size=10, color=C_INK, fill=C_TEAL_LIGHT, center=True)
        table.rows[i].cells[1].width = Cm(4.5)
        table.rows[i].cells[0].width = Cm(12)

    add_rtl_paragraph(
        doc,
        "نسخه فارسی رسمی — همراه با جداول، نمودارها و پیوست فنی",
        size=10, color=C_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=18,
    )
    add_rtl_paragraph(
        doc,
        "قلم وزیرمتن (Vazirmatn) · چینش راست‌به‌چپ · محرمانگی داده بیمار",
        size=10, bold=True, color=C_TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6,
    )
    doc.add_page_break()


def build_document():
    ASSETS.mkdir(parents=True, exist_ok=True)
    package_fonts()

    banner = ASSETS / "cover_banner.png"
    c_timeline = ASSETS / "chart_timeline.png"
    c_cost = ASSETS / "chart_cost.png"
    c_stack = ASSETS / "chart_stack.png"
    c_privacy = ASSETS / "chart_privacy.png"

    make_cover_banner(banner)
    chart_timeline(c_timeline)
    chart_cost_compare(c_cost)
    chart_stack_flow(c_stack)
    chart_privacy_scores(c_privacy)

    # cleanup tests
    for p in ASSETS.glob("_test*.png"):
        p.unlink(missing_ok=True)
    (ASSETS / "_font_test.png").unlink(missing_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), FONT_NAME)

    section = doc.sections[0]
    set_section_rtl(section)
    add_page_number_footer(section)

    build_cover(doc, banner)

    # TOC
    add_heading_rtl(doc, "فهرست مطالب", 1)
    for item in [
        "۱. درک سند و الزامات طرح",
        "۲. مشخصات کلی طرح (فرم رسمی)",
        "۳. تعریف مسأله و اهداف",
        "۴. مسیر اجرایی پیشنهادی",
        "۵. پشته فناوری فشرده و نمودار معماری",
        "۶. روش آموزش مدل‌ها به‌صورت محلی",
        "۷. مقایسه آموزش محلی و آنلاین",
        "۸. ذخیره‌سازی و محرمانگی داده",
        "۹. شیوه اشتراک‌گذاری با بیمارستان‌ها",
        "۱۰. زمان‌بندی ۱۲ ماهه",
        "۱۱. چالش‌ها، دستاوردها و تجاری‌سازی",
        "۱۲. جمع‌بندی و توصیه نهایی",
    ]:
        add_rtl_paragraph(doc, item, size=11, color=C_INK, space_after=4)
    doc.add_page_break()

    # 1
    add_heading_rtl(doc, "۱. درک سند و الزامات طرح", 1)
    add_rtl_paragraph(
        doc,
        "سند پیوست‌شده یک «فرم درخواست طرح توسعه فناوری» است؛ نه صرفاً یک ایده پژوهشی. "
        "هدف، ساخت سامانه پشتیبان تصمیم‌گیری بالینی (CDS) برای غربالگری سرطان دهانه رحم "
        "بر پایه تصاویر تمام‌اسلاید پاپ اسمیر و اطلاعات بالینی بیماران ایرانی است.",
        first_line_indent=0.5,
    )
    add_table(
        doc,
        ["الزام سند", "برداشت اجرایی"],
        [
            ["تحلیل مستقیم WSI", "بدون برش دستی سلول؛ تایل‌بندی خودکار"],
            ["طبقه‌بندی نرمال/غیرنرمال", "مدل بینایی برای هسته/سلول"],
            ["ادغام داده بالینی", "سن، سابقه درمان، علائم در تصمیم نهایی"],
            ["مجموعه‌داده بومی", "نمونه‌های واقعی ایرانی + استانداردسازی"],
            ["ارزیابی دقت و حساسیت", "شاخص‌های کمی + پایلوت بالینی"],
            ["نمونه محصول + دادگان", "prototype قابل استقرار + دیتاست نسخه‌دار"],
            ["چالش محرمانگی", "آموزش و استنتاج کاملاً محلی"],
        ],
        col_widths=[8.5, 7.5],
    )

    # 2
    add_heading_rtl(doc, "۲. مشخصات کلی طرح (فرم رسمی)", 1)
    add_heading_rtl(doc, "عنوان طرح", 2)
    add_rtl_paragraph(
        doc,
        "توسعه سامانه هوشمند شناسایی و طبقه‌بندی ناهنجاری‌های سیتولوژیک در نمونه‌های پاپ اسمیر "
        "به‌منظور تشخیص سرطان دهانه رحم در بانوان با کمک هوش مصنوعی",
        bold=True,
    )
    add_heading_rtl(doc, "خلاصه دقیق طرح", 2)
    add_rtl_paragraph(
        doc,
        "سرطان دهانه رحم از شایع‌ترین سرطان‌های قابل پیشگیری در زنان است و تشخیص زودهنگام "
        "نقش تعیین‌کننده‌ای در کاهش مرگ‌ومیر و هزینه درمان دارد. تفسیر پاپ اسمیر همچنان وابسته "
        "به نیروی متخصص است و با خطای انسانی، زمان‌بر بودن، کمبود پاتولوژیست و ناهمگونی تشخیص "
        "همراه است. این طرح یک سامانه هوشمند بومی می‌سازد که تصاویر تمام‌اسلاید (WSI) را مستقیم "
        "تحلیل می‌کند، هسته‌ها را نرمال/غیرنرمال طبقه‌بندی می‌کند و اطلاعات بالینی را در تصمیم "
        "ادغام می‌نماید. آموزش مدل‌ها روی زیرساخت داخلی انجام می‌شود تا داده بیمار به سرویس "
        "ابری خارجی ارسال نشود. خروجی، سامانه پشتیبان تصمیم بالینی فشرده و قابل نگهداری است.",
        first_line_indent=0.5,
    )
    add_table(
        doc,
        ["فیلد", "مقدار"],
        [
            ["کلمات کلیدی", "پاپ اسمیر دیجیتال، هوش مصنوعی، یادگیری ماشین، سرطان دهانه رحم، پاتولوژی دیجیتال"],
            ["مدت اجرا", "۱۲ ماه"],
            ["حوزه فناوری", "تجهیزات پزشکی / سلامت + ICT"],
            ["کاربرد نتایج", "تولید نمونه محصول + تهیه دادگان"],
            ["شیوه دانش فنی", "توسعه و تحقیق داخلی"],
            ["بازار هدف", "ملی و بین‌المللی"],
        ],
        col_widths=[11, 5],
    )

    # 3
    add_heading_rtl(doc, "۳. تعریف مسأله و اهداف", 1)
    add_rtl_paragraph(
        doc,
        "مسأله اصلی، نبود سامانه بومی، چندوجهی و قابل‌اتکاست که بتواند تصویر پاپ اسمیر را با "
        "داده بالینی ترکیب کند، روی زیرساخت محلی آموزش ببیند و در مراکز درمانی به‌صورت عملیاتی "
        "به‌کار رود. سامانه‌های خارجی معمولاً داده‌محور عمومی‌اند، فاقد بافت بالینی بومی‌اند و "
        "وابستگی ابری ایجاد می‌کنند.",
        first_line_indent=0.5,
    )
    add_heading_rtl(doc, "اهداف پنج‌گانه طرح", 2)
    for i, g in enumerate([
        "طراحی و پیاده‌سازی الگوریتم‌های هوشمند برای تحلیل تصاویر پاپ اسمیر دیجیتال",
        "توسعه مدل‌های طبقه‌بندی سلول‌های نرمال و غیرنرمال دهانه رحم",
        "ادغام داده‌های تصویری و اطلاعات بالینی بیماران در فرآیند تصمیم‌گیری",
        "ایجاد مجموعه‌داده بومی استاندارد از تصاویر پاپ اسمیر و داده‌های بالینی مرتبط",
        "ارزیابی عملکرد سامانه از نظر دقت، حساسیت و قابلیت استفاده بالینی",
    ], 1):
        add_rtl_paragraph(doc, f"{to_persian_digits(i)}. {g}", space_after=3)

    # 4
    add_heading_rtl(doc, "۴. مسیر اجرایی پیشنهادی", 1)
    add_rtl_paragraph(doc, "مسیر پیشنهادی کوتاه، قابل نگهداری و منطبق با الزامات سند است:")
    for i, s in enumerate([
        "انعقاد تفاهم با بیمارستان/دانشگاه علوم پزشکی و اخذ مجوز اخلاق",
        "ایجاد مخزن داده محلی و پروتکل بی‌نام‌سازی",
        "برچسب‌گذاری توسط پاتولوژیست در QuPath و ساخت دیتاست بومی",
        "آموزش مدل‌های تقسیم‌بندی، طبقه‌بندی و ادغام بالینی روی GPU داخلی",
        "ثبت آزمایش‌ها در MLflow و انتخاب بهترین آستانه حساسیت",
        "استقرار نمونه CDS با FastAPI روی شبکه داخلی بیمارستان",
        "پایلوت بالینی، بازخورد، مستندسازی و بسته تجاری‌سازی",
    ], 1):
        add_rtl_paragraph(doc, f"{to_persian_digits(i)}) {s}", space_after=3)

    # 5
    add_heading_rtl(doc, "۵. پشته فناوری فشرده و نمودار معماری", 1)
    add_rtl_paragraph(
        doc,
        "برای جلوگیری از پیچیدگی و هزینه نگهداری، پشته به پنج جزء اصلی کاهش یافته است. "
        "هر جزء یک مسئولیت مشخص دارد و کل جریان داخل شبکه خصوصی باقی می‌ماند.",
        first_line_indent=0.5,
    )
    add_table(
        doc,
        ["جزء", "نقش", "ورودی", "خروجی"],
        [
            ["OpenSlide", "خواندن WSI و تایل‌بندی", "فایل اسکنر", "تایل + مختصات"],
            ["QuPath", "برچسب‌گذاری پاتولوژیست", "WSI", "ماسک/برچسب"],
            ["PyTorch", "Seg + Class + Fusion", "تایل + بالینی", "مدل و امتیاز"],
            ["MLflow", "ثبت آزمایش و مدل", "متریک و وزن", "رجیستری مدل"],
            ["FastAPI", "نمونه CDS محلی", "کیس بیمار", "پیشنهاد + UI"],
        ],
        col_widths=[4.5, 4.5, 3.5, 3.5],
    )
    add_rtl_paragraph(doc, "نمودار جریان پشته فناوری:", bold=True, space_before=4)
    add_picture_centered(doc, c_stack, width_cm=16.2)
    add_heading_rtl(doc, "جزئیات لایه‌ها", 2)
    add_rtl_paragraph(
        doc,
        "لایه داده: اسکنر لام دیجیتال، NAS رمزنگاری‌شده، جدول بالینی بی‌نام. "
        "لایه یادگیری: آموزش روی یک ایستگاه GPU (مانند RTX ۴۰۹۰). "
        "لایه خدمت: FastAPI روی LAN برای نمایش سلول‌های مشکوک و امتیاز خطر؛ تصمیم نهایی با پاتولوژیست. "
        "موارد حذف‌شده برای فشرده‌سازی: Roboflow ابری برای داده بیمار، ابزارهای برچسب متعدد موازی، "
        "و MLOps سنگین در نسخه اول.",
        first_line_indent=0.5,
    )

    # 6
    add_heading_rtl(doc, "۶. روش آموزش مدل‌ها به‌صورت محلی", 1)
    add_table(
        doc,
        ["مرحله", "مدل", "برچسب لازم", "خروجی"],
        [
            ["۱. تقسیم‌بندی", "U-Net / مدل Seg", "مرز هسته/سلول", "ماسک سلول"],
            ["۲. طبقه‌بندی", "CNN/ViT", "نرمال / غیرنرمال", "احتمال سلول"],
            ["۳. ادغام بالینی", "Fusion MLP", "سن، سابقه، علائم", "خطر کیس"],
            ["۴. اختیاری MIL", "Attention-MIL", "برچسب اسلاید", "امتیاز اسلاید"],
        ],
        col_widths=[3.5, 4, 4.5, 4],
    )
    add_rtl_paragraph(doc, "اصول آموزش محلی:", bold=True)
    for t in [
        "تقسیم داده بر اساس شناسه بیمار برای جلوگیری از نشت اطلاعات",
        "افزایش داده (چرخش، تغییر رنگ رنگ‌آمیزی، تاری) برای مقاومت به تنوع اسکنر",
        "اولویت به حساسیت (Sensitivity) در غربالگری و تنظیم آستانه با پاتولوژیست",
        "ثبت همه اجراها در MLflow و ارتقای مدل فقط پس از تست نگه‌داشته‌شده",
        "بسته‌بندی مدل برای استنتاج روی همان شبکه داخلی",
    ]:
        add_rtl_paragraph(doc, f"• {t}", space_after=3)

    # 7
    add_heading_rtl(doc, "۷. مقایسه آموزش محلی و آنلاین", 1)
    add_rtl_paragraph(
        doc,
        "پلتفرم‌های آنلاین مانند Roboflow برای آزمایش‌های عمومی مفیدند، اما برای تصاویر و "
        "اطلاعات بالینی بیماران ایرانی مناسب نیستند. جدول و نمودار زیر دلیل انتخاب آموزش محلی را نشان می‌دهد.",
        first_line_indent=0.5,
    )
    add_table(
        doc,
        ["معیار", "پشته محلی", "آنلاین (مثل Roboflow)"],
        [
            ["حریم خصوصی", "داده در LAN می‌ماند", "ابر فروشنده؛ نسخه رایگان عمومی است"],
            ["تناسب با WSI و بالینی", "کامل با PyTorch سفارشی", "ضعیف برای چندوجهی/WSI"],
            ["هزینه", "GPU یک‌باره", "اشتراک ماهانه + اعتبار"],
            ["تکرار آزمایش", "نامحدود", "حدود ۳۰ دقیقه GPU ≈ ۱ اعتبار"],
            ["حاکمیت داده", "کنترل کامل", "وابستگی به SaaS خارجی"],
            ["پذیرش IT بیمارستانی", "بالا / قابل ایزوله", "اغلب مسدود برای PHI"],
        ],
        col_widths=[4, 6, 6],
    )
    add_rtl_paragraph(doc, "مقایسه کیفی تناسب:", bold=True, space_before=4)
    add_picture_centered(doc, c_privacy, width_cm=15.5)
    add_rtl_paragraph(doc, "مقایسه هزینه تقریبی سال اول:", bold=True)
    add_picture_centered(doc, c_cost, width_cm=14.5)
    add_rtl_paragraph(
        doc,
        "نتیجه: برای این طرح، آموزش و استنتاج محلی هم از نظر محرمانگی، هم هزینه بلندمدت و هم "
        "تناسب فنی برگزیده است. استفاده از سرویس آنلاین فقط برای دمو روی داده عمومی مجاز است.",
        first_line_indent=0.5,
    )

    # 8
    add_heading_rtl(doc, "۸. ذخیره‌سازی و محرمانگی داده", 1)
    add_table(
        doc,
        ["سطل داده", "محتوا", "قالب"],
        [
            ["raw_wsi", "اسلایدهای خام اسکنر", "svs / ndpi / tiff"],
            ["tiles", "قطعات تصویری + مختصات", "png/jpg + parquet"],
            ["annotations", "ماسک و چندضلعی هسته", "GeoJSON / mask"],
            ["clinical", "فیلدهای بالینی بی‌نام", "parquet / csv"],
            ["models", "وزن مدل و نسخه‌ها", "MLflow registry"],
            ["audit", "لاگ مشاهده و تأیید", "لاگ الحاق‌شونده"],
        ],
        col_widths=[3.5, 7, 5.5],
    )
    add_rtl_paragraph(doc, "قواعد حریم خصوصی:", bold=True)
    for t in [
        "عدم نگهداری نام، کد ملی و تلفن در پوشه آموزش",
        "استفاده از شناسه برگشت‌ناپذیر مطالعه (مانند CASE_۰۰۰۱۲۳)",
        "نگهداری نگاشت شناسایی روی دیسک جدا و دسترسی‌محدود",
        "رمزنگاری حجم NAS و کنترل نقش (پاتولوژیست / یادگیری ماشین / مدیر)",
        "ممنوعیت آپلود PHI به ابزار برچسب ابری یا مربی‌های SaaS",
        "پشتیبان روزانه محلی + نسخه سرد هفتگی",
    ]:
        add_rtl_paragraph(doc, f"• {t}", space_after=3)

    # 9
    add_heading_rtl(doc, "۹. شیوه اشتراک‌گذاری با بیمارستان‌ها", 1)
    add_rtl_paragraph(
        doc,
        "اصل راهنما: نرم‌افزار، پروتکل، وزن مدل و بسته ارزیابی بی‌نام قابل اشتراک است؛ "
        "بایگانی خامِ قابل شناسایی بیماران قابل اشتراک عمومی نیست.",
        first_line_indent=0.5,
    )
    add_table(
        doc,
        ["حالت همکاری", "آنچه بیمارستان دریافت می‌کند", "آنچه بازمی‌گرداند"],
        [
            ["شریک داده", "پروتکل QuPath و قالب پروژه", "WSI و بالینی بی‌نام (دیسک امن/VPN)"],
            ["سایت پایلوت", "بسته CDS روی LAN", "بازخورد پاتولوژیست و متریک استفاده"],
            ["به‌روزرسانی مدل", "وزن مدل جدید", "شاخص‌های تجمیعی اختیاری"],
            ["چندمرکزی (آتی)", "اسکیما و SOP مشترک", "خروجی بی‌نام هماهنگ"],
        ],
        col_widths=[3.5, 6.5, 6],
    )
    add_rtl_paragraph(doc, "پیش‌نیازهای حکمرانی قبل از هر اشتراک:", bold=True)
    for t in [
        "تأیید اخلاق / IRB",
        "توافق‌نامه اشتراک داده (DSA)",
        "راستی‌آزمایی بی‌نام‌سازی",
        "تعریف نقش‌های دسترسی و تماس پاسخ به رخداد",
    ]:
        add_rtl_paragraph(doc, f"• {t}", space_after=3)

    # 10
    add_heading_rtl(doc, "۱۰. زمان‌بندی ۱۲ ماهه", 1)
    add_picture_centered(doc, c_timeline, width_cm=16.2)
    add_table(
        doc,
        ["فاز", "ماه", "خروجی کلیدی"],
        [
            ["P0 راه‌اندازی", "۱–۲", "مجوز اخلاق، تفاهم‌نامه، سخت‌افزار و NAS"],
            ["P1 داده", "۲–۵", "دیتاست بومی نسخه ۱"],
            ["P2 مدل", "۴–۸", "مدل‌های Seg/Class/Fusion با متریک قابل قبول"],
            ["P3 سامانه", "۷–۱۰", "نمونه FastAPI قابل استقرار روی LAN"],
            ["P4 پایلوت", "۹–۱۱", "گزارش قابلیت استفاده بالینی"],
            ["P5 اختتام", "۱۱–۱۲", "مقاله، کارت دیتاست، بسته تجاری‌سازی"],
        ],
        col_widths=[4, 3, 9],
    )
    add_rtl_paragraph(doc, "تناظر اهداف فرم با ماه‌های کلیدی:", bold=True)
    add_table(
        doc,
        ["هدف فرم", "ماه تحقق"],
        [
            ["الگوریتم تحلیل تصویر", "ماه ۶"],
            ["مدل نرمال/غیرنرمال", "ماه ۸"],
            ["ادغام بالینی", "ماه ۹"],
            ["مجموعه‌داده بومی", "ماه ۵ (v1) و ماه ۱۱ (v1.1)"],
            ["ارزیابی دقت/حساسیت/کاربری", "ماه‌های ۱۰ تا ۱۲"],
        ],
        col_widths=[9, 7],
    )

    # 11
    add_heading_rtl(doc, "۱۱. چالش‌ها، دستاوردها و تجاری‌سازی", 1)
    add_table(
        doc,
        ["چالش", "راهکار"],
        [
            ["کمبود دادگان بومی", "همکاری با دانشگاه علوم پزشکی و چند مرکز"],
            ["تنوع کیفیت لام", "پروتکل استاندارد + افزایش داده + مدل مقاوم"],
            ["محرمانگی بیمار", "بی‌نام‌سازی + آموزش محلی + کنترل دسترسی"],
            ["عدم توازن کلاس", "وزن‌دهی خطا و آستانه حساسیت‌محور"],
            ["ریسک پیچیدگی نرم‌افزار", "قانون پنج‌ابزاری پشته فشرده"],
        ],
        col_widths=[6, 10],
    )
    add_heading_rtl(doc, "سایر دستاوردها", 2)
    for t in [
        "انتشار مقالات علمی در مجلات و کنفرانس‌های معتبر",
        "ایجاد مجموعه‌داده بومی قابل استفاده در پژوهش‌های آینده",
        "توسعه نمونه اولیه سامانه تشخیص هوشمند با استقرار محلی",
        "زمینه‌سازی برای ثبت اختراع یا تجاری‌سازی محصول",
        "ارتقاء کیفیت و هوشمندسازی فرآیندهای تشخیصی موجود در کشور",
    ]:
        add_rtl_paragraph(doc, f"• {t}", space_after=3)

    add_heading_rtl(doc, "آمادگی تجاری‌سازی", 2)
    add_table(
        doc,
        ["عنوان", "پاسخ"],
        [
            ["دستاورد نهایی", "تولید فناوری"],
            ["تمایل به تجاری‌سازی", "توسط تیم مجری"],
            ["بهره‌بردار اصلی", "تمامی بانوان (جمعیت غربالگری)"],
            ["نهادهای استفاده‌کننده", "بیمارستان‌ها و مراکز درمانی دولتی و خصوصی"],
            ["حمایت مورد نیاز", "همکاری دانشگاه علوم پزشکی استان برای جمع‌آوری داده"],
            ["گام پس از اتمام", "تجاری‌سازی برای فراگیری در کشور"],
        ],
        col_widths=[5.5, 10.5],
    )

    # 12
    add_heading_rtl(doc, "۱۲. جمع‌بندی و توصیه نهایی", 1)
    add_rtl_paragraph(
        doc,
        "این پیشنهادیه دقیقاً در قالب فرم توسعه فناوری تنظیم شده و مسیر فنی روشنی ارائه می‌دهد: "
        "پشته فشرده محلی، آموزش روی GPU داخلی، ذخیره امن داده بی‌نام، اشتراک کنترل‌شده با بیمارستان‌ها، "
        "و زمان‌بندی ۱۲ ماهه منطبق با پنج هدف سند. این مسیر ضمن حفظ محرمانگی، قابلیت نگهداری بلندمدت "
        "و بومی‌سازی فناوری سلامت دیجیتال را فراهم می‌کند.",
        first_line_indent=0.5,
    )
    add_rtl_paragraph(doc, "جریان نهایی پیشنهادی:", bold=True, space_before=8)
    add_rtl_paragraph(
        doc,
        "WSI محلی ← OpenSlide ← برچسب QuPath ← آموزش PyTorch روی GPU داخلی ← MLflow ← FastAPI روی LAN ← تأیید پاتولوژیست",
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=C_TEAL_DARK, size=10,
    )
    add_horizontal_line(doc, C_GOLD)
    add_rtl_paragraph(
        doc,
        "پایان پیشنهادیه — نسخه فارسی با قلم وزیرمتن (Vazirmatn) و چینش راست‌به‌چپ",
        size=9, color=C_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10,
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Wrote {OUT_DOCX}")
    return OUT_DOCX


if __name__ == "__main__":
    build_document()
